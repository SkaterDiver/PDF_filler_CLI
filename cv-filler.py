#!/usr/bin/env python3
"""Cover Letter CLI - Fill templates with job-specific information."""

import os
import re
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document

TEMPLATES_DIR = Path(__file__).parent / "Templates"
OUTPUTS_DIR = Path(__file__).parent / "Outputs"
PLACEHOLDER_PATTERN = re.compile(r"\[([^\]]+)\]")

# LibreOffice paths for Windows
SOFFICE_PATHS = [
    Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
    Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
]


def get_soffice_path():
    """Find LibreOffice soffice executable."""
    for path in SOFFICE_PATHS:
        if path.exists():
            return str(path)
    return "soffice"  # Fall back to PATH


def list_templates():
    """List all .docx templates in the Templates folder."""
    templates = list(TEMPLATES_DIR.glob("*.docx"))
    return sorted(templates)


def extract_placeholders_from_text(text):
    """Extract all [placeholder] patterns from text."""
    return set(PLACEHOLDER_PATTERN.findall(text))


def extract_placeholders(doc):
    """Extract all unique placeholders from a document."""
    placeholders = set()

    # Check paragraphs
    for para in doc.paragraphs:
        placeholders.update(extract_placeholders_from_text(para.text))

    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    placeholders.update(extract_placeholders_from_text(para.text))

    return sorted(placeholders)


def prompt_for_values(placeholders):
    """Prompt user for each placeholder value."""
    values = {}
    print("\nEnter values for each field:")
    print("-" * 40)
    for placeholder in placeholders:
        if placeholder.lower() == "date":
            value = datetime.now().strftime("%Y-%m-%d")
            print(f"  {placeholder}: {value} (auto)")
        else:
            value = input(f"  {placeholder}: ").strip()
        values[placeholder] = value
    return values


def replace_in_paragraph(paragraph, values):
    """Replace placeholders in a paragraph while preserving formatting."""
    # Get the full text to check for placeholders
    full_text = paragraph.text

    # Check if any placeholders exist in this paragraph
    if not PLACEHOLDER_PATTERN.search(full_text):
        return

    # Replace in each run
    for run in paragraph.runs:
        for placeholder, value in values.items():
            bracketed = f"[{placeholder}]"
            if bracketed in run.text:
                run.text = run.text.replace(bracketed, value)


def fill_template(doc, values):
    """Replace all placeholders in the document with provided values."""
    # Process paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para, values)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, values)


def sanitize_filename(name):
    """Remove invalid characters from filename."""
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        name = name.replace(char, "")
    return name.strip()


def save_as_pdf(doc, company_name):
    """Save the document as PDF using LibreOffice."""
    # Create output filename
    date_str = datetime.now().strftime("%Y-%m-%d")
    safe_company = sanitize_filename(company_name) if company_name else "Unknown"
    pdf_name = f"CoverLetter_{safe_company}_{date_str}.pdf"

    # Save temp docx file
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        temp_docx = tmp.name

    try:
        doc.save(temp_docx)

        # Convert to PDF using LibreOffice
        result = subprocess.run(
            [
                get_soffice_path(),
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(OUTPUTS_DIR),
                temp_docx
            ],
            capture_output=True,
            text=True
        )

        if result.returncode != 0:
            print(f"\nError converting to PDF: {result.stderr}")
            return None

        # Rename the output file
        temp_pdf_name = Path(temp_docx).stem + ".pdf"
        temp_pdf_path = OUTPUTS_DIR / temp_pdf_name
        final_pdf_path = OUTPUTS_DIR / pdf_name

        # Handle existing file with same name
        if final_pdf_path.exists():
            counter = 1
            while final_pdf_path.exists():
                pdf_name = f"CoverLetter_{safe_company}_{date_str}_{counter}.pdf"
                final_pdf_path = OUTPUTS_DIR / pdf_name
                counter += 1

        if temp_pdf_path.exists():
            temp_pdf_path.rename(final_pdf_path)
            return final_pdf_path
        else:
            print("\nError: PDF was not created.")
            return None

    finally:
        # Clean up temp file
        if os.path.exists(temp_docx):
            os.remove(temp_docx)


def main():
    """Main CLI loop."""
    print("\n" + "=" * 50)
    print("       Cover Letter Generator")
    print("=" * 50)

    while True:
        # List templates
        templates = list_templates()

        if not templates:
            print("\nNo templates found in Templates folder.")
            break

        print("\nAvailable templates:")
        print("-" * 40)
        for i, template in enumerate(templates, 1):
            # Clean up template name for display
            name = template.stem.replace("[Template]_", "").replace("_", " ")
            print(f"  {i}. {name}")
        print(f"\n  Type 'exit' to quit")

        # Get user selection
        choice = input("\nSelect template number: ").strip().lower()

        if choice == "exit":
            print("\nGoodbye!")
            break

        try:
            idx = int(choice) - 1
            if 0 <= idx < len(templates):
                template_path = templates[idx]
            else:
                print("\nInvalid selection. Please try again.")
                continue
        except ValueError:
            print("\nInvalid input. Please enter a number or 'exit'.")
            continue

        # Load template
        print(f"\nLoading template: {template_path.name}")
        doc = Document(template_path)

        # Extract placeholders
        placeholders = extract_placeholders(doc)

        if not placeholders:
            print("\nNo placeholders found in this template.")
            continue

        print(f"\nFound {len(placeholders)} placeholder(s):")
        for p in placeholders:
            print(f"  - [{p}]")

        # Get values from user
        values = prompt_for_values(placeholders)

        # Fill template
        fill_template(doc, values)

        # Determine company name for filename
        company_name = values.get("Company Name") or values.get("Company") or values.get("Employer") or ""

        # Save as PDF
        print("\nGenerating PDF...")
        pdf_path = save_as_pdf(doc, company_name)

        if pdf_path:
            print(f"\nSaved: {pdf_path}")

        print("\n" + "-" * 50)


if __name__ == "__main__":
    main()
